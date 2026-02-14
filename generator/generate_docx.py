#!/usr/bin/env python3
"""
Convert book.json to book.docx using python-docx.
Styling optimized for Kindle Create / Vellum.

Usage:
    python generator/generate_docx.py outputs/decision-making/book.json
"""

import sys
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_content(content: str, descriptor: str) -> str:
    """
    Strips the initial H1 (Name) and Italicized Descriptor line from the content.
    Expects format:
    # Name
    
    *Descriptor*
    
    Body text...
    """
    if not content:
        return ""
        
    lines = content.split('\n')
    
    # Helper to check if line is empty
    def is_empty(idx):
        return idx < len(lines) and not lines[idx].strip()
        
    idx = 0
    # 1. Skip leading empty lines
    while is_empty(idx):
        idx += 1
        
    # 2. Skip H1 (Name)
    if idx < len(lines) and lines[idx].strip().startswith('# '):
        idx += 1
        
    # 3. Skip lines between H1 and Descriptor
    while is_empty(idx):
        idx += 1
        
    # 4. Skip Descriptor
    if descriptor and idx < len(lines):
        stripped = lines[idx].strip()
        # Check for strict match of *Descriptor*
        if stripped == f"*{descriptor}*":
             idx += 1
             
    # 5. Skip remaining empty lines (until body)
    while is_empty(idx):
        idx += 1
        
    return '\n'.join(lines[idx:])

def add_markdown_content(doc, content: str):
    """
    Parses simple markdown content and adds it to the document.
    Handles:
    - Headers (#, ##) -> Headings
    - Scene breaks (***) -> Centered separator
    - Bold (**text**)
    - Italic (*text*)
    - Custom Subtitles (*Text but not ended with *) -> Centered Italic
    - Section Headers (> Section:) -> Centered + Spacing
    """
    if not content:
        return

    lines = content.split('\n')
    title_processed = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
            
        # Skip Footer
        if stripped.startswith('Part of') and 'alexandria.press' in stripped:
            continue

        # Handle repeated Title (Content starting with H1)
        # We assume the first line being H1 is a repetition of the chapter title
        if not title_processed:
            if stripped.startswith('# '):
                title_processed = True
                continue
            else:
                title_processed = True
            
        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
            
        # Scene Breaks
        if stripped in ['***', '* * *', '---']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('***')
            run.bold = True
            continue
            
        # Custom Subtitle: Starts with * but not * (bullet) or ** (bold) or *** (break)
        # Note: We check if it DOES NOT end with * to distinguish from standard italics
        if stripped.startswith('*') and not stripped.startswith('**') and not stripped.startswith('* ') and not stripped.endswith('*'):
             p = doc.add_paragraph()
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             # Strip the leading *
             text = stripped[1:].strip()
             run = p.add_run(text)
             run.italic = True
             continue
             
        # Section Header - Remove entirely
        # Matches "> Section:", "> **Section:**", "Section:", etc.
        clean_line = stripped.replace('*', '').replace('>', '').replace('**', '').strip()
        if clean_line.lower().startswith('section:'):
            continue

        # Bullet List (Lines starting with "- " or "* ")
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:].strip()
            format_and_add_run(p, text)
            continue

        # Numbered List (Lines starting with "1. ", "2. ", etc.)
        # Simple check for "digits. "
        match_number = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if match_number:
            p = doc.add_paragraph(style='List Number')
            text = match_number.group(2).strip()
            format_and_add_run(p, text)
            continue
            
        # Normal Paragraph
        p = doc.add_paragraph()
        format_and_add_run(p, stripped)

def format_and_add_run(paragraph, text):
    """
    Parses **bold** and *italic* and adds runs to paragraph.
    Handles nested or adjacent formatting by splitting tokens.
    """
    # Define a simple tokenizer for bold and italic
    # This regex looks for **...** or *...*
    # It captures the delimiters to help identify what matched
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for token in tokens:
        if not token:
            continue
            
        if token.startswith('**') and token.endswith('**'):
            # Bold
            content = token[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            # Italic
            content = token[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(token)

def generate_docx(book_path: str):
    """Generate DOCX from book.json"""
    path = Path(book_path)
    if not path.exists():
        print(f"Error: File not found: {path}")
        sys.exit(1)
        
    print(f"Loading {path}...")
    with open(path, 'r', encoding='utf-8') as f:
        book = json.load(f)
        
    doc = Document()
    
    # --- Title Page ---
    for _ in range(4): doc.add_paragraph() # Spacing
    
    title = doc.add_paragraph(book.get('title', 'Untitled'))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = 'Title'
    
    if book.get('descriptor'):
        desc = doc.add_paragraph(book['descriptor'])
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc.style = 'Subtitle'
        
    doc.add_paragraph("\n\n")
    pub = doc.add_paragraph("Alexandria Press")
    pub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # --- Introduction ---
    if book.get('introduction'):
        doc.add_heading('Introduction', level=1)
        add_markdown_content(doc, book['introduction'])
        doc.add_page_break()
        
    # --- Entries ---
    entries = book.get('entries', [])
    print(f"Processing {len(entries)} entries...")
    
    for i, entry in enumerate(entries):
        # Chapter Title (Heading 1)
        doc.add_heading(entry.get('name', f"Chapter {i+1}"), level=1)
        
        # Descriptor (Heading 2)
        if entry.get('descriptor'):
            doc.add_heading(entry['descriptor'], level=2)
        
        # Content
        content = entry.get('content', '')
        content = clean_content(content, entry.get('descriptor'))
        add_markdown_content(doc, content)
        
        # Page break after every chapter except the last one
        if i < len(entries) - 1:
            doc.add_page_break()
            
    # --- Save ---
    output_path = path.with_suffix('.docx')
    doc.save(output_path)
    
    print(f"\n✓ Generated: {output_path}")
    print(f"  Size: {output_path.stat().st_size / 1024:.1f} KB")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python generator/generate_docx.py <path/to/book.json>")
        sys.exit(1)
    
    generate_docx(sys.argv[1])