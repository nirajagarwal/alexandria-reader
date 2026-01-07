#!/usr/bin/env python3
"""
Export book data from outputs folder to Microsoft Word documents.

This script reads book.json files from the outputs directory and generates
professionally formatted Word documents (.docx) for each book.
"""

import json
import os
import sys
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx library not found.")
    print("Install it with: pip install python-docx")
    sys.exit(1)


def create_styled_document() -> Document:
    """Create a new Word document."""
    doc = Document()
    return doc


def add_book_header(doc: Document, book_data: Dict[str, Any]) -> None:
    """Add book title, descriptor, and introduction to the document."""
    # Title
    title_para = doc.add_paragraph(book_data.get('title', 'Untitled'))
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(12)
    
    # Descriptor
    if 'descriptor' in book_data:
        desc_para = doc.add_paragraph(book_data['descriptor'])
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.runs[0]
        desc_run.font.name = 'Georgia'
        desc_run.font.size = Pt(14)
        desc_run.italic = True
        desc_run.font.color.rgb = RGBColor(80, 80, 80)
        desc_para.paragraph_format.space_after = Pt(24)
    
    # Introduction
    if 'introduction' in book_data:
        intro_heading = doc.add_heading('Introduction', level=1)
        intro_para = doc.add_paragraph(book_data['introduction'])
        intro_para.paragraph_format.space_after = Pt(24)
    
    # Add page break after header
    doc.add_page_break()


def parse_markdown_content(content: str) -> list:
    """
    Parse markdown content into structured blocks.
    Returns a list of tuples: (type, content)
    Types: 'heading1', 'heading2', 'heading3', 'paragraph', 'blockquote'
    """
    blocks = []
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Heading 1
        if line.startswith('# '):
            blocks.append(('heading1', line[2:].strip()))
        # Heading 2
        elif line.startswith('## '):
            blocks.append(('heading2', line[3:].strip()))
        # Heading 3
        elif line.startswith('### '):
            blocks.append(('heading3', line[4:].strip()))
        # Blockquote
        elif line.startswith('> '):
            quote_lines = [line[2:]]
            i += 1
            while i < len(lines) and lines[i].startswith('> '):
                quote_lines.append(lines[i][2:])
                i += 1
            blocks.append(('blockquote', ' '.join(quote_lines)))
            continue
        # Horizontal rule
        elif line.strip() in ['---', '***', '___']:
            blocks.append(('hr', ''))
        # Paragraph (including bold/italic markdown)
        else:
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('>'):
                para_lines.append(lines[i])
                i += 1
            blocks.append(('paragraph', ' '.join(para_lines)))
            continue
        
        i += 1
    
    return blocks


def apply_inline_formatting(paragraph, text: str) -> None:
    """Apply bold and italic formatting to text within a paragraph."""
    import re
    
    # Pattern to match ***text***, **text**, or *text*
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('***') and part.endswith('***'):
            # Bold and italic
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Normal text
            paragraph.add_run(part)


def add_entry_content(doc: Document, entry: Dict[str, Any]) -> None:
    """Add a single entry's content to the document."""
    content = entry.get('content', '')
    
    # Parse markdown content
    blocks = parse_markdown_content(content)
    
    for block_type, block_content in blocks:
        if block_type == 'heading1':
            # Skip the main title as we'll use entry name
            continue
        elif block_type == 'heading2':
            heading = doc.add_heading(block_content, level=2)
        elif block_type == 'heading3':
            heading = doc.add_heading(block_content, level=3)
        elif block_type == 'blockquote':
            para = doc.add_paragraph()
            apply_inline_formatting(para, block_content)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
            para_format = para.paragraph_format
            para_format.space_before = Pt(6)
            para_format.space_after = Pt(6)
            # Make it italic and gray
            for run in para.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
        elif block_type == 'hr':
            # Add a horizontal line (using a bottom border on an empty paragraph)
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
        elif block_type == 'paragraph':
            para = doc.add_paragraph()
            apply_inline_formatting(para, block_content)


def add_entry(doc: Document, entry: Dict[str, Any], book_id: str) -> None:
    """Add a complete entry to the document."""
    # Entry title
    entry_name = entry.get('name', entry.get('title', 'Untitled Entry'))
    
    # For periodic-tales, include symbol and atomic number
    if book_id == 'periodic-tales':
        symbol = entry.get('symbol', '')
        atomic_number = entry.get('atomic_number', '')
        if symbol and atomic_number:
            title_text = f"{entry_name} ({symbol}, #{atomic_number})"
        else:
            title_text = entry_name
    else:
        title_text = entry_name
    
    # Add entry title with formatting
    title_para = doc.add_paragraph(title_text)
    title_run = title_para.runs[0]
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(12)
    
    # Entry descriptor if available
    if 'descriptor' in entry:
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(entry['descriptor'])
        desc_run.italic = True
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = RGBColor(80, 80, 80)
        desc_para.paragraph_format.space_after = Pt(12)
    
    # Entry content
    add_entry_content(doc, entry)
    
    # Add page break after each entry (except the last one)
    doc.add_page_break()


def export_book_to_word(book_path: Path, output_dir: Path) -> None:
    """Export a single book to a Word document."""
    print(f"Processing: {book_path}")
    
    # Read book.json
    try:
        with open(book_path, 'r', encoding='utf-8') as f:
            book_data = json.load(f)
    except Exception as e:
        print(f"Error reading {book_path}: {e}")
        return
    
    # Create document
    doc = create_styled_document()
    
    # Add book header
    add_book_header(doc, book_data)
    
    # Add entries
    book_id = book_data.get('book_id', book_path.parent.name)
    entries = book_data.get('entries', [])
    
    print(f"  Found {len(entries)} entries")
    
    for i, entry in enumerate(entries, 1):
        print(f"  Processing entry {i}/{len(entries)}: {entry.get('name', entry.get('title', 'Untitled'))}")
        add_entry(doc, entry, book_id)
    
    # Remove the last page break
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        if last_para.text == '':
            # Remove empty paragraph (page break)
            p = last_para._element
            p.getparent().remove(p)
    
    # Save document
    book_title = book_data.get('title', 'Book').replace(' ', '_').replace(':', '')
    output_file = output_dir / f"{book_id}.docx"
    
    try:
        doc.save(str(output_file))
        print(f"✓ Saved: {output_file}")
    except Exception as e:
        print(f"Error saving {output_file}: {e}")


def main():
    """Main function to process all books in the outputs folder."""
    # Get the project root directory
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    outputs_dir = project_root / 'outputs'
    
    # Create output directory for Word documents
    word_output_dir = project_root / 'outputs_word'
    word_output_dir.mkdir(exist_ok=True)
    
    print(f"Alexandria Press - Book to Word Exporter")
    print(f"=" * 50)
    print(f"Outputs directory: {outputs_dir}")
    print(f"Word output directory: {word_output_dir}")
    print()
    
    # Find all book.json files
    book_files = list(outputs_dir.glob('*/book.json'))
    
    if not book_files:
        print("No book.json files found in outputs directory.")
        return
    
    print(f"Found {len(book_files)} books to process")
    print()
    
    # Process each book
    for book_file in sorted(book_files):
        export_book_to_word(book_file, word_output_dir)
        print()
    
    print("=" * 50)
    print(f"✓ Export complete! Word documents saved to: {word_output_dir}")


if __name__ == '__main__':
    main()
